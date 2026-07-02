from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from rembg import remove
import cv2
import numpy as np
import io
import os
from contextlib import asynccontextmanager

# Load Super Resolution models on startup
sr_models = {}

@asynccontextmanager
async def lifespan(app: FastAPI):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    models_dir = os.path.join(base_dir, "models")
    
    # Try to load FSRCNN x2
    model_x2_path = os.path.join(models_dir, "FSRCNN_x2.pb")
    if os.path.exists(model_x2_path):
        sr_x2 = cv2.dnn_superres.DnnSuperResImpl_create()
        sr_x2.readModel(model_x2_path)
        sr_x2.setModel("fsrcnn", 2)
        sr_models[2] = sr_x2
        print("Loaded FSRCNN_x2")
        
    # Try to load FSRCNN x4
    model_x4_path = os.path.join(models_dir, "FSRCNN_x4.pb")
    if os.path.exists(model_x4_path):
        sr_x4 = cv2.dnn_superres.DnnSuperResImpl_create()
        sr_x4.readModel(model_x4_path)
        sr_x4.setModel("fsrcnn", 4)
        sr_models[4] = sr_x4
        print("Loaded FSRCNN_x4")
    yield

app = FastAPI(lifespan=lifespan)

# Izinkan Frontend React mengakses Backend (CORS)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/remove-bg")
async def remove_bg(file: UploadFile = File(...)):
    try:
        input_data = await file.read()
        output_data = remove(input_data)
        return Response(content=output_data, media_type="image/png")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upscale")
async def upscale_image(scale: int = Form(2), file: UploadFile = File(...)):
    if scale not in sr_models:
        raise HTTPException(status_code=400, detail=f"Scale x{scale} model not loaded or unsupported.")
    
    try:
        input_data = await file.read()
        nparr = np.frombuffer(input_data, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            raise HTTPException(status_code=400, detail="Invalid image format")
            
        result = sr_models[scale].upsample(img)
        
        _, encoded_img = cv2.imencode('.png', result)
        output_data = encoded_img.tobytes()
        
        return Response(content=output_data, media_type="image/png")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/auto-enhance")
async def auto_enhance(file: UploadFile = File(...)):
    try:
        input_data = await file.read()
        nparr = np.frombuffer(input_data, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            raise HTTPException(status_code=400, detail="Invalid image format")
        
        # 1. CLAHE for dynamic range and contrast balance
        lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
        l, a, b = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        cl = clahe.apply(l)
        limg = cv2.merge((cl, a, b))
        enhanced_contrast = cv2.cvtColor(limg, cv2.COLOR_LAB2BGR)
        
        # 2. Bilateral Filter for edge-preserving noise reduction
        denoised = cv2.bilateralFilter(enhanced_contrast, d=5, sigmaColor=35, sigmaSpace=35)
        
        # 3. Unsharp Masking for clean and natural sharpening
        blurred = cv2.GaussianBlur(denoised, (5, 5), 1.0)
        sharpened = cv2.addWeighted(denoised, 1.5, blurred, -0.5, 0)
        
        # 4. Color Saturation Boost (convert to HSV and increase S channel gently)
        hsv = cv2.cvtColor(sharpened, cv2.COLOR_BGR2HSV)
        h, s, v = cv2.split(hsv)
        s = np.clip(s * 1.15, 0, 255).astype(np.uint8)
        enhanced_hsv = cv2.merge((h, s, v))
        result = cv2.cvtColor(enhanced_hsv, cv2.COLOR_HSV2BGR)
        
        _, encoded_img = cv2.imencode('.png', result)
        output_data = encoded_img.tobytes()
        
        return Response(content=output_data, media_type="image/png")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/denoise")
async def apply_denoise(strength: float = Form(10.0), file: UploadFile = File(...)):
    try:
        input_data = await file.read()
        nparr = np.frombuffer(input_data, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if img is None:
            raise HTTPException(status_code=400, detail="Invalid image format")
            
        # Denoising
        result = cv2.fastNlMeansDenoisingColored(img, None, h=strength, hColor=strength, templateWindowSize=7, searchWindowSize=21)
        
        _, encoded_img = cv2.imencode('.png', result)
        output_data = encoded_img.tobytes()
        
        return Response(content=output_data, media_type="image/png")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ==========================================
# DOCUMENT PROCESSING MODULE
# ==========================================
import fitz  # PyMuPDF
import zipfile
import shutil
import tempfile
import sys
from pdf2docx import Converter
from docx import Document

# COM Context dependency for Windows Office COM
def com_context():
    if sys.platform == "win32":
        import comtypes
        comtypes.CoInitializeEx(comtypes.COINIT_APARTMENTTHREADED)
        try:
            yield
        finally:
            comtypes.CoUninitialize()
    else:
        yield

def convert_office_to_pdf_win(input_path: str, output_path: str, ext: str):
    import comtypes.client
    
    if ext in [".docx", ".doc"]:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            doc = word.Documents.Open(input_path)
            doc.SaveAs(output_path, FileFormat=17) # 17 is wdFormatPDF
            doc.Close()
        finally:
            word.Quit()
            
    elif ext in [".xlsx", ".xls"]:
        excel = comtypes.client.CreateObject('Excel.Application')
        excel.Visible = False
        excel.DisplayAlerts = False
        try:
            wb = excel.Workbooks.Open(input_path)
            wb.ExportAsFixedFormat(0, output_path) # 0 is xlTypePDF
            wb.Close(False)
        finally:
            excel.Quit()
            
    elif ext in [".pptx", ".ppt"]:
        powerpoint = comtypes.client.CreateObject('PowerPoint.Application')
        powerpoint.DisplayAlerts = 1 # ppAlertsNone (or 1 depending on version, generally 1 is ppAlertsNone in early versions, 7 is ppAlertsNone in later)
        try:
            pres = powerpoint.Presentations.Open(input_path, WithWindow=False)
            pres.SaveAs(output_path, 32) # 32 is ppSaveAsPDF
            pres.Close()
        finally:
            powerpoint.Quit()
    else:
        raise ValueError(f"Extension {ext} is not supported for Office to PDF.")

def compress_pdf_images(doc, quality: int, dpi_target: int):
    from PIL import Image as PILImage
    import io as pillow_io
    
    processed_xrefs = set()
    
    # Determine maximum dimension based on target DPI
    max_dim = 1200
    if dpi_target == 100:
        max_dim = 800
    elif dpi_target == 150:
        max_dim = 1200
    elif dpi_target == 200:
        max_dim = 1600
        
    for page in doc:
        try:
            img_list = page.get_images(full=True)
        except Exception:
            continue
            
        for img in img_list:
            xref = img[0]
            if xref in processed_xrefs:
                continue
            processed_xrefs.add(xref)
            
            try:
                # Get the original stream size to check if compression actually helps
                try:
                    original_stream_len = len(doc.xref_stream(xref))
                except Exception:
                    original_stream_len = 0
                
                pix = fitz.Pixmap(doc, xref)
                
                # Convert to RGB if not Gray or RGB
                if pix.colorspace.n not in (1, 3):
                    try:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    except Exception:
                        continue
                    
                width = pix.width
                height = pix.height
                if width <= 0 or height <= 0:
                    continue
                    
                # Handle alpha channel safely
                has_alpha = pix.alpha or pix.colorspace.n == 4
                if has_alpha:
                    try:
                        img_data = PILImage.frombytes("RGBA", [width, height], pix.samples)
                        background = PILImage.new("RGB", img_data.size, (255, 255, 255))
                        background.paste(img_data, mask=img_data.split()[3])
                        img_data = background
                    except Exception:
                        # Fallback if alpha pasting fails
                        img_data = PILImage.frombytes("RGB", [width, height], pix.samples[:width*height*3])
                else:
                    mode = "L" if pix.colorspace.n == 1 else "RGB"
                    img_data = PILImage.frombytes(mode, [width, height], pix.samples)
                
                # Resize if the image exceeds the max dimension
                if width > max_dim or height > max_dim:
                    scale = max_dim / max(width, height)
                    new_w = int(width * scale)
                    new_h = int(height * scale)
                    img_data = img_data.resize((new_w, new_h), PILImage.Resampling.LANCZOS)
                    
                # Compress image to JPEG
                buffer = pillow_io.BytesIO()
                img_data.save(buffer, format="JPEG", quality=quality)
                compressed_bytes = buffer.getvalue()
                
                # Replace only if the compressed size is smaller
                if original_stream_len == 0 or len(compressed_bytes) < original_stream_len:
                    page.replace_image(xref, stream=compressed_bytes)
            except Exception as e:
                print(f"Failed to compress image xref {xref}: {e}")

@app.post("/compress-pdf")
async def compress_pdf(level: str = Form("medium"), file: UploadFile = File(...)):
    try:
        file_bytes = await file.read()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        # Scrub metadata safely
        try:
            doc.scrub(metadata=True, xml_metadata=True, thumbnails=True, reset_fields=True)
        except Exception:
            pass
            
        # Determine image compression quality and DPI target
        quality = 65
        dpi_target = 150
        if level == "low":
            quality = 85
            dpi_target = 200
        elif level == "medium":
            quality = 65
            dpi_target = 150
        elif level == "high":
            quality = 40
            dpi_target = 100
            
        # Compress images manually
        compress_pdf_images(doc, quality, dpi_target)
            
        # garbage=3 is fast and cleans unused objects
        output_bytes = doc.write(garbage=3, deflate=True)
        doc.close()
        
        return Response(content=output_bytes, media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal mengompresi PDF: {str(e)}")

@app.post("/merge-pdf")
def merge_pdf(files: list[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Minimal unggah 2 file PDF untuk digabungkan.")
        
    try:
        merged_doc = fitz.open()
        opened_docs = []
        
        for f in files:
            f_bytes = f.file.read()
            doc = fitz.open(stream=f_bytes, filetype="pdf")
            merged_doc.insert_pdf(doc)
            opened_docs.append(doc)
            
        output_bytes = merged_doc.write()
        merged_doc.close()
        for doc in opened_docs:
            doc.close()
            
        return Response(content=output_bytes, media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal menggabungkan PDF: {str(e)}")

@app.post("/split-pdf")
def split_pdf(ranges: str = Form(...), file: UploadFile = File(...)):
    try:
        file_bytes = file.file.read()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        total = len(doc)
        
        new_doc = fitz.open()
        
        for part in ranges.split(','):
            part = part.strip()
            if not part:
                continue
            if '-' in part:
                start_str, end_str = part.split('-')
                start = max(1, min(total, int(start_str)))
                end = max(1, min(total, int(end_str)))
                if start <= end:
                    for p in range(start, end + 1):
                        new_doc.insert_pdf(doc, from_page=p-1, to_page=p-1)
                else:
                    for p in range(start, end - 1, -1):
                        new_doc.insert_pdf(doc, from_page=p-1, to_page=p-1)
            else:
                p = max(1, min(total, int(part)))
                new_doc.insert_pdf(doc, from_page=p-1, to_page=p-1)
                
        output_bytes = new_doc.write()
        new_doc.close()
        doc.close()
        
        return Response(content=output_bytes, media_type="application/pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal memisahkan PDF: {str(e)}")

@app.post("/extract-pdf-images")
def extract_pdf_images(file: UploadFile = File(...)):
    try:
        file_bytes = file.file.read()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                zf.writestr(f"page_{i+1}.png", img_bytes)
                
        output_bytes = zip_buffer.getvalue()
        doc.close()
        
        return Response(
            content=output_bytes, 
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=extracted_pages.zip"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal mengekstrak gambar dari PDF: {str(e)}")

@app.post("/pdf-to-word")
def pdf_to_word(ocr: bool = Form(False), file: UploadFile = File(...)):
    temp_dir = tempfile.mkdtemp()
    input_path = os.path.join(temp_dir, "input.pdf")
    output_path = os.path.join(temp_dir, "output.docx")
    
    try:
        file_bytes = file.file.read()
        with open(input_path, "wb") as f:
            f.write(file_bytes)
            
        if ocr:
            try:
                import pytesseract
                doc = fitz.open(input_path)
                docx_doc = Document()
                
                for i, page in enumerate(doc):
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes("png")
                    
                    from PIL import Image as PILImage
                    import io as pillow_io
                    img = PILImage.open(pillow_io.BytesIO(img_data))
                    
                    text = pytesseract.image_to_string(img)
                    docx_doc.add_heading(f"Halaman {i+1}", level=2)
                    docx_doc.add_paragraph(text)
                    docx_doc.add_page_break()
                    
                docx_doc.save(output_path)
                doc.close()
            except Exception as ocr_err:
                print("OCR library error, falling back to layout extraction:", ocr_err)
                doc = fitz.open(input_path)
                docx_doc = Document()
                for i, page in enumerate(doc):
                    text = page.get_text()
                    docx_doc.add_heading(f"Halaman {i+1}", level=2)
                    docx_doc.add_paragraph(text)
                    docx_doc.add_page_break()
                docx_doc.save(output_path)
                doc.close()
        else:
            cv = Converter(input_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
            
        with open(output_path, "rb") as f:
            docx_bytes = f.read()
            
        return Response(
            content=docx_bytes, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        print("PDF to Word conversion failed:", e)
        raise HTTPException(status_code=500, detail=f"Gagal mengonversi PDF ke Word: {str(e)}")
        
    finally:
        try:
            shutil.rmtree(temp_dir)
        except Exception:
            pass

@app.post("/office-to-pdf")
def office_to_pdf(file: UploadFile = File(...)):
    if sys.platform != "win32":
        raise HTTPException(
            status_code=400, 
            detail="Fitur konversi Office ke PDF hanya didukung di OS Windows dengan Microsoft Office terpasang secara lokal."
        )
        
    import comtypes
    comtypes.CoInitializeEx(comtypes.COINIT_APARTMENTTHREADED)
    
    try:
        filename = file.filename
        _, ext = os.path.splitext(filename.lower())
        
        if ext not in [".docx", ".doc", ".xlsx", ".xls", ".pptx", ".ppt"]:
            raise HTTPException(
                status_code=400, 
                detail="Format file tidak didukung. Harap upload file Word, Excel, atau PowerPoint."
            )
            
        temp_dir = tempfile.mkdtemp()
        input_path = os.path.join(temp_dir, f"input{ext}")
        output_path = os.path.join(temp_dir, "output.pdf")
        
        try:
            with open(input_path, "wb") as f:
                f.write(file.file.read())
                
            convert_office_to_pdf_win(input_path, output_path, ext)
            
            with open(output_path, "rb") as f:
                pdf_bytes = f.read()
                
            return Response(content=pdf_bytes, media_type="application/pdf")
            
        except Exception as e:
            print("Office to PDF Error:", e)
            raise HTTPException(status_code=500, detail=f"Gagal mengonversi file ke PDF: {str(e)}")
            
        finally:
            try:
                shutil.rmtree(temp_dir)
            except Exception:
                pass
    finally:
        comtypes.CoUninitialize()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)